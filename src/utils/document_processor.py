"""Document processing utilities for multiple file formats."""
import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import time
import warnings
import re
import logging

# PDF processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX processing imports
try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Suppress warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pdfminer")
warnings.filterwarnings("ignore", category=DeprecationWarning, module="pdfminer")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handle document processing for multiple file formats (PDF, DOCX, TXT)."""
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 200, use_cache: bool = True):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_cache = use_cache

        # Optimized text splitter for compliance documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            # Enhanced separators for better compliance document chunking
            separators=[
                "\n\n\n",      # Multiple line breaks (section separators)
                "\n\n",        # Paragraph breaks
                "\n",          # Line breaks
                ". ",          # Sentence endings with space
                "! ",          # Exclamation with space
                "? ",          # Question with space
                "; ",          # Semicolon with space
                ", ",          # Comma with space
                " ",           # Space
                ""             # Character level
            ]
        )

        # Compliance-specific text splitter for regulatory documents
        self.compliance_splitter = RecursiveCharacterTextSplitter(
            chunk_size=600,  # Smaller chunks for better precision
            chunk_overlap=150,
            length_function=len,
            separators=[
                "\nArticle ",   # Legal articles
                "\nSection ",  # Legal sections
                "\n(",         # Numbered items
                "\na) ",       # Lettered items
                "\nb) ",
                "\nc) ",
                "\n1. ",       # Numbered lists
                "\n2. ",
                "\n3. ",
                "\n\n",        # Paragraph breaks
                "\n",          # Line breaks
                ". ",          # Sentence endings
                " ",           # Space
                ""             # Character level
            ]
        )
        
        # Check available dependencies
        self._check_dependencies()

    def _choose_optimal_splitter(self, text_content: str) -> RecursiveCharacterTextSplitter:
        """Choose the optimal text splitter based on document content."""
        # Check for compliance/legal document indicators
        compliance_indicators = [
            'article ', 'section ', 'regulation ', 'requirement', 'obligation',
            'gdpr', 'compliance', 'framework', 'standard', 'procedure',
            'a) ', 'b) ', 'c) ', '1. ', '2. ', '3. '
        ]

        text_lower = text_content.lower()
        compliance_score = sum(1 for indicator in compliance_indicators if indicator in text_lower)

        # Use compliance splitter for documents with high compliance content
        if compliance_score >= 3:
            logger.info("Using compliance-optimized text splitter")
            return self.compliance_splitter
        else:
            logger.info("Using standard text splitter")
            return self.text_splitter
    
    def _check_dependencies(self):
        """Check which document processing libraries are available."""
        if not PYMUPDF_AVAILABLE and not PYPDF2_AVAILABLE:
            logger.warning("No PDF processing libraries available. Install PyMuPDF or PyPDF2.")
        
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX files won't be processed.")
    
    def process_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process a document of any supported format."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type by extension
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._process_pdf(file_path, metadata)
        elif extension == '.docx':
            return self._process_docx(file_path, metadata)
        elif extension == '.txt':
            return self._process_txt(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _process_pdf(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process PDF files."""
        if PYMUPDF_AVAILABLE:
            return self._process_pdf_pymupdf(file_path, metadata)
        elif PYPDF2_AVAILABLE:
            return self._process_pdf_pypdf2(file_path, metadata)
        else:
            raise ImportError("No PDF processing library available. Install PyMuPDF or PyPDF2.")
    
    def _process_pdf_pymupdf(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process PDF using PyMuPDF (fastest option)."""
        try:
            doc = fitz.open(str(file_path))
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
                text_content += "\n\n"  # Add page separator
            
            doc.close()
            
            # Choose optimal splitter and split into chunks
            optimal_splitter = self._choose_optimal_splitter(text_content)
            chunks = optimal_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    'source': str(file_path),
                    'file_type': 'pdf',
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'processing_method': 'pymupdf'
                }
                if metadata:
                    doc_metadata.update(metadata)
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Processed PDF {file_path.name}: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF with PyMuPDF: {e}")
            raise
    
    def _process_pdf_pypdf2(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process PDF using PyPDF2 (fallback option)."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text
                    text_content += "\n\n"  # Add page separator
            
            # Choose optimal splitter and split into chunks
            optimal_splitter = self._choose_optimal_splitter(text_content)
            chunks = optimal_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    'source': str(file_path),
                    'file_type': 'pdf',
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'processing_method': 'pypdf2'
                }
                if metadata:
                    doc_metadata.update(metadata)
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Processed PDF {file_path.name}: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF with PyPDF2: {e}")
            raise
    
    def _process_docx(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process DOCX files."""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(str(file_path))
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            # Choose optimal splitter and split into chunks
            optimal_splitter = self._choose_optimal_splitter(text_content)
            chunks = optimal_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    'source': str(file_path),
                    'file_type': 'docx',
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'processing_method': 'python-docx'
                }
                if metadata:
                    doc_metadata.update(metadata)
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Processed DOCX {file_path.name}: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise
    
    def _process_txt(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process TXT files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise ValueError(f"Could not decode text file {file_path} with any of the attempted encodings")
            
            # Choose optimal splitter and split into chunks
            optimal_splitter = self._choose_optimal_splitter(text_content)
            chunks = optimal_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    'source': str(file_path),
                    'file_type': 'txt',
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'processing_method': 'text_file'
                }
                if metadata:
                    doc_metadata.update(metadata)
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Processed TXT {file_path.name}: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of a file."""
        sha256_hash = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                # Read file in chunks to handle large files efficiently
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash for {file_path}: {e}")
            return ""
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file metadata including hash."""
        file_path = Path(file_path)
        try:
            stat = file_path.stat()
            file_hash = self.calculate_file_hash(str(file_path))
            
            return {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_size": stat.st_size,
                "file_hash": file_hash,
                "modified_time": stat.st_mtime,
                "created_time": stat.st_ctime,
                "file_extension": file_path.suffix.lower()
            }
        except Exception as e:
            logger.error(f"Error getting file metadata for {file_path}: {e}")
            return {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_extension": file_path.suffix.lower(),
                "error": str(e)
            }
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        extensions = ['.txt']  # Always supported
        
        if PYMUPDF_AVAILABLE or PYPDF2_AVAILABLE:
            extensions.append('.pdf')
        
        if PYTHON_DOCX_AVAILABLE:
            extensions.append('.docx')
        
        return extensions
    
    def is_supported_file(self, file_path: str) -> bool:
        """Check if file type is supported."""
        extension = Path(file_path).suffix.lower()
        return extension in self.get_supported_extensions()
